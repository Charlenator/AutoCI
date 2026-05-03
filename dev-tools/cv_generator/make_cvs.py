"""Convert LLM-generated CV JSON batches into .docx files.

Usage:
    cd dev-tools/cv_generator
    python make_cvs.py cvs_batch_1.json
    python make_cvs.py *.json    # process every batch in this folder

Outputs go to ./output/<filename>.docx (filename pulled from each CV's "filename" field,
with .txt swapped to .docx).

Requires:
    pip install python-docx

The script is intentionally forgiving — it strips markdown fences if the LLM
included them, validates each CV row independently, and skips/logs malformed
entries rather than aborting the whole batch.
"""

from __future__ import annotations

import argparse
import glob
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    print("ERROR: python-docx is not installed.")
    print("  Run: pip install python-docx")
    sys.exit(1)


def strip_md_fences(raw: str) -> str:
    """LLMs sometimes wrap JSON in ```json ... ```. Strip those if present."""
    text = raw.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*\n", "", text)
        text = re.sub(r"\n```\s*$", "", text)
    return text


def docx_filename(cv: dict) -> str:
    raw = cv.get("filename", "").strip()
    if not raw:
        # Fallback: name + role + index hash
        name = cv.get("name", "unknown").lower().replace(" ", "_")
        role = cv.get("role_target", "x").lower().replace(" ", "_")
        raw = f"{role}_{name}.docx"
    if raw.endswith(".txt"):
        raw = raw[:-4] + ".docx"
    if not raw.endswith(".docx"):
        raw = raw + ".docx"
    return raw


def render_cv(cv: dict, out_path: Path) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    name_p = doc.add_paragraph()
    name_run = name_p.add_run(cv.get("name", "Unknown Candidate"))
    name_run.bold = True
    name_run.font.size = Pt(20)

    contact_bits = [
        cv.get("email", "").strip(),
        cv.get("phone", "").strip(),
        cv.get("location", "").strip(),
    ]
    contact_line = " | ".join(b for b in contact_bits if b)
    if contact_line:
        cp = doc.add_paragraph()
        cp.add_run(contact_line).font.size = Pt(10)

    target = cv.get("role_target")
    yoe = cv.get("years_experience")
    if target or yoe is not None:
        tp = doc.add_paragraph()
        bits = []
        if target:
            bits.append(f"Target role: {target}")
        if yoe is not None:
            bits.append(f"{yoe} years experience")
        run = tp.add_run("  •  ".join(bits))
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    summary = cv.get("summary", "").strip()
    if summary:
        doc.add_heading("Summary", level=2)
        for para in summary.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    skills = cv.get("skills") or []
    if skills:
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(str(s) for s in skills))

    experience = cv.get("experience") or []
    if experience:
        doc.add_heading("Experience", level=2)
        for job in experience:
            company = job.get("company", "")
            title = job.get("title", "")
            start = job.get("start_year", "")
            end = job.get("end_year", "")
            header_p = doc.add_paragraph()
            header_run = header_p.add_run(f"{title} — {company}")
            header_run.bold = True
            header_p.add_run(f"  ({start} – {end})")
            for hl in job.get("highlights", []) or []:
                bullet = doc.add_paragraph(str(hl), style="List Bullet")
                bullet.paragraph_format.left_indent = Pt(18)

    education = cv.get("education") or []
    if education:
        doc.add_heading("Education", level=2)
        for ed in education:
            line = f"{ed.get('degree', '')} — {ed.get('institution', '')} ({ed.get('year', '')})"
            doc.add_paragraph(line.strip())

    cur = cv.get("current_salary_zar")
    exp = cv.get("expected_salary_zar")
    if cur or exp:
        doc.add_heading("Compensation", level=2)
        bits = []
        if cur:
            bits.append(f"Current: R{cur:,}")
        if exp:
            bits.append(f"Expected: R{exp:,}")
        doc.add_paragraph(" | ".join(bits))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def process_file(json_path: Path, out_dir: Path) -> tuple[int, int]:
    raw = json_path.read_text(encoding="utf-8")
    cleaned = strip_md_fences(raw)
    try:
        cvs = json.loads(cleaned)
    except json.JSONDecodeError as exc:
        print(f"  [{json_path.name}] JSON parse error: {exc}")
        return (0, 1)

    if not isinstance(cvs, list):
        print(f"  [{json_path.name}] expected a JSON array, got {type(cvs).__name__}")
        return (0, 1)

    written = 0
    failed = 0
    for i, cv in enumerate(cvs):
        if not isinstance(cv, dict):
            print(f"  [{json_path.name}] entry {i}: not an object, skipping")
            failed += 1
            continue
        try:
            target = out_dir / docx_filename(cv)
            render_cv(cv, target)
            print(f"    -> {target.name}")
            written += 1
        except Exception as exc:
            print(f"  [{json_path.name}] entry {i} ({cv.get('name', '?')}): {exc}")
            failed += 1
    return (written, failed)


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert CV JSON batches into .docx files.")
    parser.add_argument("paths", nargs="+", help="JSON files (globs allowed, e.g. *.json)")
    parser.add_argument("--out", default="output", help="Output directory (default: ./output)")
    args = parser.parse_args()

    here = Path(__file__).parent
    out_dir = Path(args.out)
    if not out_dir.is_absolute():
        out_dir = here / out_dir

    expanded: list[Path] = []
    for raw in args.paths:
        matches = [Path(p) for p in glob.glob(raw)] if any(c in raw for c in "*?[") else [Path(raw)]
        expanded.extend(matches)

    if not expanded:
        print("No JSON files matched.")
        sys.exit(1)

    total_written = 0
    total_failed = 0
    for p in expanded:
        if not p.exists():
            print(f"  [skip] {p}: not found")
            continue
        print(f"Processing {p} ...")
        w, f = process_file(p, out_dir)
        total_written += w
        total_failed += f

    print()
    print(f"Done. {total_written} .docx written, {total_failed} failed.")
    # Some Windows consoles use cp1252 and choke on emoji in the absolute path.
    # Encode-safely so the script doesn't end on a scary traceback.
    out_str = str(out_dir.resolve())
    try:
        print(f"Output: {out_str}")
    except UnicodeEncodeError:
        print(f"Output: {out_str.encode('ascii', errors='replace').decode('ascii')}")
        print("       (path contained non-ascii chars; consoles in cp1252 mode can't print the originals)")


if __name__ == "__main__":
    main()
