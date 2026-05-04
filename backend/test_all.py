#!/usr/bin/env python3
"""AutoCI — Multi-level test suite. Run: python3 test_all.py [--level 1|2|3|4|5] [--filter "B7:"]

Acronym legend for test names (alphabetical):
  B-aug     Live-search augmentation (Tavily/News/Adzuna → corpus_chunks)
  B-evidence Source-record evidence for aggregate SQL results
  B6        Resend send API wrapper (resend_client.py)
  B7        cal.com slot-lookup wrapper (cal_com_client.py)
  01.4      CV section-based smart-chunking (cv_chunking.py)
  01.5      Email vectorizer (email_vectorizer.py)
  D1/D2/D3  Detection agents (internal benchmarks, external, gap analysis)
  K1–K7     Kaizen phase agents (Define/Measure/Analyse/5 Whys/Ishikawa/Improve/Control)
  O3        Phase-gate enforcer
  S1        Query Planner Agent (LLM-based plan construction)
  S3        SQL Executor (template + freeform with allowlist)
  S5        CV classifier agent (is-this-a-CV)
  S6        .docx field extractor agent
  S7        Confidentiality classifier agent
  T1        AnalyticsLibrary (MCP-aligned analytical tools)
  T2        ValidationInterceptor (Pydantic schema validation wrapper)

Use --filter to run only tests whose name contains the given string. Example:
  python test_all.py --level 1 --filter "B7:"    # only B7 CalComClient tests
  python test_all.py --filter "S6:"              # only S6 CV extractor tests (level 1+)
"""

import sys
import json
import os
from datetime import date

# ── CLI args ──────────────────────────────────────────────────────────────

_FILTER: str | None = None
_LEVEL: int = 0
_TARGET: str = "local"  # "local" or "modal"

argv = sys.argv[1:]
i = 0
while i < len(argv):
    if argv[i] == "--filter" and i + 1 < len(argv):
        _FILTER = argv[i + 1].lower()
        i += 2
    elif argv[i] == "--level" and i + 1 < len(argv):
        _LEVEL = int(argv[i + 1])
        i += 2
    elif argv[i] == "--target" and i + 1 < len(argv):
        _TARGET = argv[i + 1].lower()
        if _TARGET not in ("local", "modal"):
            print(f"Unknown target '{_TARGET}'. Use 'local' or 'modal'.")
            sys.exit(1)
        i += 2
    else:
        i += 1

BACKEND_BASE = (
    "https://charlenator--autoci-backend-fastapi-app.modal.run"
    if _TARGET == "modal"
    else "http://localhost:8000"
)

# ── Level 1: Unit Tests (0 dependencies) ──────────────────────────────────

ANSI = {"GREEN": "\033[92m", "RED": "\033[91m", "YELLOW": "\033[93m", "RESET": "\033[0m"}
pass_count = 0
fail_count = 0

def test(name, fn):
    """Register a test. If --filter is active, skip unless name contains it.

    The callable must return a truthy value or raise. Returning ``False``
    or ``None`` is treated as a failure — this catches tests that forget
    to assert (e.g. ``test("x", lambda: False)``).
    """
    global pass_count, fail_count
    if _FILTER and _FILTER not in name.lower():
        return  # skip silently
    try:
        result = fn()
        if not result:
            raise AssertionError(f"returned {result}")
        print(f"  {ANSI['GREEN']}✓{ANSI['RESET']} {name}")
        pass_count += 1
    except Exception as e:
        print(f"  {ANSI['RED']}✗{ANSI['RESET']} {name}: {e}")
        fail_count += 1

def level1_unit():
    print(f"\n{'='*60}\n📦 LEVEL 1: Unit Tests\n{'='*60}")

    # T1 — AnalyticsLibrary
    from api.tools.t1_mcp_analytics import AnalyticsLibrary
    a = AnalyticsLibrary()

    test("TTF with empty data", lambda: a.time_to_fill([], []) == 0.0)

    events = [
        {"candidate_id": "A", "stage": "Applied", "event_date": date(2025,1,1), "outcome": "Advanced"},
        {"candidate_id": "A", "stage": "Offer", "event_date": date(2025,2,1), "outcome": "Offer Extended"},
    ]
    hires = [{"candidate_id": "A", "accepted": True}]
    ttf = a.time_to_fill(events, hires)
    test(f"TTF single hire (expected 31): got {ttf}", lambda: abs(ttf - 31.0) < 1)

    test("Conversion 100%", lambda: abs(a.stage_conversion_rate(events, "Applied") - 1.0) < 0.01)
    test("Dropoff 0%", lambda: abs(a.stage_dropoff_rate(events, "Applied") - 0.0) < 0.01)

    oar = a.offer_acceptance_rate(hires, [{"outcome": "Accepted"}])
    test("OAR 100%", lambda: abs(oar - 1.0) < 0.01)

    bench = a.benchmark_comparison(45.0, 35.0)
    test("Benchmark ~28.6% delta", lambda: abs(bench["delta_pct"] - 28.6) < 1)

    outliers = a.outlier_detection([10, 20, 15, 100, 12, 18])
    test(f"Outlier detection: index 3 (value 100) = {outliers}", lambda: 3 in outliers)

    # T2 — ValidationInterceptor
    from api.tools.t2_validation_interceptor import validate_agent_output
    from pydantic import BaseModel

    class TestSchema(BaseModel):
        value: float

    @validate_agent_output(schema=TestSchema, min_sample_size=3)
    def good_output():
        return {"value": 42.0}

    @validate_agent_output(schema=TestSchema, min_sample_size=3)
    def bad_output_shape():
        return {"wrong_key": 1}

    r1, v1 = good_output()
    test("T2: valid output passes", lambda: v1.passed)

    r2, v2 = bad_output_shape()
    test("T2: bad schema fails", lambda: not v2.passed)

    # S1 — QueryPlannerAgent (Sprint B1, replaces the keyword-based TranslationAgent)
    from api.agents.specialists.s1_query_planner import QueryPlannerAgent, QueryPlan
    from api.agents.specialists.sql_templates import (
        TEMPLATE_REGISTRY,
        build_template_sql,
        build_template_evidence_sql,
        TemplateParamError,
    )
    test("S1: planner imports cleanly", lambda: QueryPlannerAgent is not None)
    test("S1: template registry has at least 6 templates",
         lambda: len(TEMPLATE_REGISTRY) >= 6)
    test("S1: every registry entry has a builder",
         lambda: all(t.build is not None for t in TEMPLATE_REGISTRY.values()))

    # Verify each template builds with sensible inputs (or throws cleanly when required params missing).
    sample_params = {
        "time_to_fill": {},
        "offer_acceptance_rate": {},
        "conversion_rate": {},
        "kpis_for_role": {"role_title": "Java"},
        "pipeline_volume_by_stage": {},
        "candidate_search_by_skill": {"skill_keyword": "Java"},
        "candidate_by_email": {"email": "test@example.com"},
        "industry_benchmark_for_role": {"role_title": "UX"},
    }
    for tid, params in sample_params.items():
        if tid not in TEMPLATE_REGISTRY:
            continue
        tid_, params_ = tid, params
        test(f"S1: template '{tid_}' builds SELECT",
             lambda tid=tid_, p=params_: build_template_sql(tid, p).strip().upper().startswith(("SELECT", "WITH")))

    # Required-param validation should raise.
    try:
        build_template_sql("kpis_for_role", {})
        kpis_required_raised = False
    except TemplateParamError:
        kpis_required_raised = True
    test("S1: missing required param raises TemplateParamError", lambda: kpis_required_raised)

    # S5 — CVClassifierAgent
    from api.agents.specialists.s5_cv_classifier import CVClassifierAgent
    from api.tools.t3_litellm_router import LiteLLMRouter
    s5 = CVClassifierAgent(LiteLLMRouter())
    test("S5: agent imports cleanly", lambda: CVClassifierAgent is not None)
    result_empty = s5.is_cv("")
    test("S5: empty string -> is_cv=False",
         lambda: result_empty["is_cv"] is False)
    test("S5: empty string -> confidence 0.0",
         lambda: result_empty["confidence"] == 0.0)
    test("S5: empty string -> reason 'empty input'",
         lambda: result_empty["reason"] == "empty input")
    result_ws = s5.is_cv("   \n  \t  ")
    test("S5: whitespace-only -> is_cv=False",
         lambda: result_ws["is_cv"] is False)
    test("S5: whitespace-only -> confidence 0.0",
         lambda: result_ws["confidence"] == 0.0)
    from api.agents.specialists.s5_cv_classifier import _parse_json
    test("S5: _parse_json handles None", lambda: _parse_json(None) is None)
    test("S5: _parse_json handles valid JSON", lambda: _parse_json('{"is_cv": true}') == {"is_cv": True})
    test("S5: _parse_json strips markdown fences", lambda: _parse_json("```json\n{\"is_cv\": false}\n```") == {"is_cv": False})

    # S6 — CVExtractorAgent
    from api.agents.specialists.s6_cv_extractor import (
        CVExtractorAgent,
        _clean_str,
        _normalize_skills,
        _normalize_record,
        _fallback_record,
    )
    test("S6: agent imports cleanly", lambda: CVExtractorAgent is not None)

    # Helper: _clean_str
    test("S6: _clean_str strips whitespace", lambda: _clean_str("  hello  ") == "hello")
    test("S6: _clean_str returns None for None", lambda: _clean_str(None) is None)
    test("S6: _clean_str returns None for empty", lambda: _clean_str("  ") is None)

    # Helper: _normalize_skills
    test("S6: _normalize_skills lowercases + dedupes",
         lambda: _normalize_skills(["Python", "python", "Java"]) == ["python", "java"])
    test("S6: _normalize_skills handles None", lambda: _normalize_skills(None) == [])
    test("S6: _normalize_skills handles non-list", lambda: _normalize_skills("Python") == ["python"])

    # Helper: _normalize_record with known data
    _rec_in = {"name": " John  ", "email": " JOHN@EXAMPLE.COM ", "phone": "+27 123 4567",
               "summary": "Experienced dev", "skills": ["Python", "python"], "experience": [], "education": []}
    _rec_out = _normalize_record(_rec_in, "raw text here")
    test("S6: _normalize_record strips name", lambda: _rec_out["name"] == "John")
    test("S6: _normalize_record lowercases email", lambda: _rec_out["email"] == "john@example.com")
    test("S6: _normalize_record dedupes skills", lambda: _rec_out["skills"] == ["python"])
    test("S6: _normalize_record no missing fields",
         lambda: _rec_out["missing_fields"] == [])
    test("S6: _normalize_record preserves raw_text", lambda: _rec_out["raw_text"] == "raw text here")

    # Helper: _normalize_record with sparse data → all missing
    _rec_sparse = _normalize_record({}, "empty")
    test("S6: sparse record has all 7 missing fields",
         lambda: set(_rec_sparse["missing_fields"]) == {"name", "email", "phone", "summary", "skills", "experience", "education"})

    # Helper: _fallback_record
    _fb = _fallback_record("some text")
    test("S6: fallback has raw_text", lambda: _fb["raw_text"] == "some text")
    test("S6: fallback has all missing_fields", lambda: len(_fb["missing_fields"]) == 7)

    # In-memory .docx test: constructs a docx with known content and checks
    # that the extractor gets raw_text populated (LLM not required since it
    # will fall back gracefully when DEEPSEEK_API_KEY is absent).
    try:
        from docx import Document
        import io
        _mem_doc = Document()
        _mem_doc.add_paragraph("John Doe")
        _mem_doc.add_paragraph("john@example.com")
        _mem_doc.add_paragraph("+27 12 345 6789")
        _mem_doc.add_paragraph("Experienced software engineer with 5 years in Python.")
        _mem_buf = io.BytesIO()
        _mem_doc.save(_mem_buf)
        _mem_buf.seek(0)
        from api.tools.t3_litellm_router import LiteLLMRouter
        _s6 = CVExtractorAgent(LiteLLMRouter())
        _s6_result = _s6.extract(_mem_buf.read())
        test("S6: in-memory docx extracts raw_text with known name",
             lambda: "John Doe" in (_s6_result.get("raw_text") or ""))
        test("S6: in-memory docx extracts raw_text with known email",
             lambda: "john@example.com" in (_s6_result.get("raw_text") or ""))
        # Without LLM key, fields will fall back; test that fallback structure is correct
        test("S6: docx extract returns dict with all expected keys",
             lambda: all(k in _s6_result for k in ("name", "email", "phone", "summary", "skills", "experience", "education", "missing_fields", "raw_text")))
    except ImportError:
        test("S6: python-docx not installed — skipping docx tests", lambda: True)

    # S7 — ConfidentialityAgent
    from api.agents.specialists.s7_confidentiality import ConfidentialityAgent
    from api.tools.t3_litellm_router import LiteLLMRouter
    s7 = ConfidentialityAgent(LiteLLMRouter())

    test("S7: agent imports cleanly", lambda: ConfidentialityAgent is not None)
    result_empty = s7.classify("")
    test("S7: empty string -> confidential=True", lambda: result_empty["confidential"] is True)
    test("S7: empty string -> reason 'empty input'", lambda: result_empty["reason"] == "empty input")
    result_ws = s7.classify("   \n  ")
    test("S7: whitespace-only -> confidential=True", lambda: result_ws["confidential"] is True)

    # S7: _parse_json helper (reused from S5, but tested through S7 path)
    test("S7: _parse_json handles valid JSON", lambda: _parse_json('{"confidential": false, "reason": "ok"}') == {"confidential": False, "reason": "ok"})
    test("S7: _parse_json handles None", lambda: _parse_json(None) is None)

    # 01.4 — cv_chunking
    from api.workers.cv_chunking import chunk_cv

    _full_rec = {
        "name": "Jane Doe",
        "email": "jane@example.com",
        "phone": "+27 82 123 4567",
        "summary": "Experienced data scientist with 6 years in ML.",
        "skills": ["python", "machine learning", "sql"],
        "experience": [
            {"role": "Senior Data Scientist", "company": "Acme Corp",
             "start_date": "2021-03", "end_date": "2025-01",
             "description": "Built ML pipelines for recommendation systems."},
            {"role": "Data Analyst", "company": "Beta Inc",
             "start_date": "2018-06", "end_date": "2021-02",
             "description": "Created dashboards and automated reporting."},
        ],
        "education": [
            {"school": "University of Cape Town", "degree": "MSc Data Science", "year": "2018"},
        ],
        "missing_fields": [],
        "raw_text": "Jane Doe\njane@example.com\n+27 82 123 4567\n...",
    }
    _full_chunks = chunk_cv(_full_rec, "cand-001")
    test("01.4: full record produces 6 chunks (identity+skills+summary+2exp+education)",
         lambda: len(_full_chunks) == 6)
    _kinds = {c["metadata"]["chunk_kind"] for c in _full_chunks}
    test("01.4: chunk_kinds cover all 5 types",
         lambda: _kinds == {"identity", "skills", "summary", "experience", "education"})
    test("01.4: identity chunk has name + email",
         lambda: "Jane Doe" in _full_chunks[0]["chunk_text"] and "jane@example.com" in _full_chunks[0]["chunk_text"])
    test("01.4: experience chunks have role_target in metadata",
         lambda: any(c["metadata"].get("role_target") == "Senior Data Scientist" for c in _full_chunks))
    test("01.4: all chunks are corpus_name='cvs'",
         lambda: all(c["corpus_name"] == "cvs" for c in _full_chunks))
    test("01.4: all chunks have confidential=True in metadata",
         lambda: all(c["metadata"]["confidential"] is True for c in _full_chunks))

    # Record with no summary → 5 chunks (summary skipped)
    _no_summary = dict(_full_rec, summary="")
    _no_summary_chunks = chunk_cv(_no_summary, "cand-002")
    test("01.4: no summary produces 5 chunks (summary skipped)",
         lambda: len(_no_summary_chunks) == 5)
    test("01.4: no summary — 'summary' not in chunk_kinds",
         lambda: "summary" not in {c["metadata"]["chunk_kind"] for c in _no_summary_chunks})

    # 01.5 — email_vectorizer
    from api.workers.email_vectorizer import vectorize_email

    _ev = vectorize_email("FW: Candidate CV", "Please find attached the CV for John.\n\nRegards, Recruiter", "inb-001")
    test("01.5: typical email produces 2 chunks (subject + body)",
         lambda: len(_ev) == 2)
    _ev_kinds = {c["metadata"]["chunk_kind"] for c in _ev}
    test("01.5: chunk kinds are subject and body",
         lambda: _ev_kinds == {"subject", "body"})
    test("01.5: subject chunk has correct text",
         lambda: any(c["chunk_text"] == "FW: Candidate CV" for c in _ev))
    test("01.5: all chunks are corpus_name='inbound_emails'",
         lambda: all(c["corpus_name"] == "inbound_emails" for c in _ev))
    test("01.5: all chunks have confidential=True",
         lambda: all(c["metadata"]["confidential"] is True for c in _ev))

    _ev_empty = vectorize_email("", "", "inb-002")
    test("01.5: empty email produces 0 chunks",
         lambda: len(_ev_empty) == 0)

    _ev_subject_only = vectorize_email("Just a subject", "", "inb-003")
    test("01.5: subject-only email produces 1 chunk",
         lambda: len(_ev_subject_only) == 1 and _ev_subject_only[0]["metadata"]["chunk_kind"] == "subject")

    # B-aug: planner envelope honours live-search fields and the sanitizer
    # auto-fills sources + forces RAG when a stripped envelope arrives.
    from api.agents.specialists.s1_query_planner import (
        QueryPlan,
        _envelope_to_plan,
        _sanitize_plan,
        _clean_live_sources,
    )
    test("B-aug: _clean_live_sources accepts known names",
         lambda: _clean_live_sources(["Tavily", "news", "garbage"]) == ["tavily", "news"])
    test("B-aug: _clean_live_sources tolerates None / non-list",
         lambda: _clean_live_sources(None) == [] and _clean_live_sources("adzuna") == ["adzuna"])
    _aug_env = {
        "needs_sql": False,
        "needs_rag": False,
        "needs_live_search": True,
        "live_search_sources": ["adzuna", "tavily"],
        "live_search_topic": "Senior Java Developer",
        "explanation": "live",
        "confidence": 0.8,
    }
    _aug_plan = _sanitize_plan(_envelope_to_plan("current java salaries", _aug_env))
    test("B-aug: live-search plan parses sources + topic",
         lambda: _aug_plan.needs_live_search and _aug_plan.live_search_sources == ["adzuna", "tavily"])
    test("B-aug: sanitizer forces needs_rag=True when live search fires",
         lambda: _aug_plan.needs_rag is True and _aug_plan.rag_query == "current java salaries")
    _aug_env_min = {"needs_live_search": True, "explanation": "x", "confidence": 0.5}
    _aug_min_plan = _sanitize_plan(_envelope_to_plan("recent SA tech news", _aug_env_min))
    test("B-aug: sanitizer auto-fills sources when planner forgets them",
         lambda: set(_aug_min_plan.live_search_sources) == {"tavily", "news", "adzuna"} and _aug_min_plan.live_search_topic == "recent SA tech news")

    # B-evidence: every template that exposes build_evidence must produce a
    # SELECT/WITH-prefixed SQL string for representative params.
    evidence_params = {
        "time_to_fill": {"role_title": "Java"},
        "offer_acceptance_rate": {"role_title": "Java"},
        "conversion_rate": {"role_title": "Java"},
        "kpis_for_role": {"role_title": "Java"},
        "pipeline_volume_by_stage": {},
    }
    for tid, params in evidence_params.items():
        if tid not in TEMPLATE_REGISTRY:
            continue
        tid_, params_ = tid, params
        test(f"B-evidence: '{tid_}' builds source-record SELECT",
             lambda tid=tid_, p=params_: (build_template_evidence_sql(tid, p) or "").strip().upper().startswith(("SELECT", "WITH")))
    # Templates whose result is already record-level should return None.
    test("B-evidence: candidate_search_by_skill has no evidence (already record-level)",
         lambda: build_template_evidence_sql("candidate_search_by_skill", {"skill_keyword": "Java"}) is None)

    # SQL Executor sanity (regex allowlist on freeform)
    from api.agents.specialists.s3_sql_executor import (
        SQLExecutor,
        FreeformSQLValidationError,
        _validate_freeform_sql,
    )
    try:
        _validate_freeform_sql("DROP TABLE candidates")
        drop_blocked = False
    except FreeformSQLValidationError:
        drop_blocked = True
    test("S3: freeform DROP rejected by allowlist", lambda: drop_blocked)
    try:
        _validate_freeform_sql("SELECT 1; SELECT 2")
        stacked_blocked = False
    except FreeformSQLValidationError:
        stacked_blocked = True
    test("S3: stacked statement rejected", lambda: stacked_blocked)
    try:
        _validate_freeform_sql("SELECT count(*) FROM candidates")
        select_ok = True
    except FreeformSQLValidationError:
        select_ok = False
    test("S3: clean SELECT passes allowlist", lambda: select_ok)

    # B6 — ResendClient (simple inline assertions)
    from api.integrations.resend_client import ResendClient, ResendError

    class _FakeResponse:
        def __init__(self, status_code, json_data):
            self.status_code = status_code
            self._json = json_data
            self.text = json.dumps(json_data)
        @property
        def is_success(self):
            return self.status_code < 400
        def json(self):
            return self._json

    # Use an explicit call-counter to track which httpx.Client instance to inspect.
    _b6_last_body = {"val": None}

    class _FakeClient:
        def __init__(self, timeout=None):
            self.timeout = timeout
        def __enter__(self):
            return self
        def __exit__(self, *args):
            pass
        def post(self, url, *, headers, json):
            _b6_last_body["val"] = json
            return _FakeResponse(200, {"id": "abc-123"})

    class _FakeErrorClient:
        def __init__(self, timeout=None):
            self.timeout = timeout
        def __enter__(self):
            return self
        def __exit__(self, *args):
            pass
        def post(self, url, *, headers, json):
            return _FakeResponse(422, {"error": "missing from field"})

    _b6_mod = __import__("api.integrations.resend_client", fromlist=["httpx"])
    _b6_orig = _b6_mod.httpx.Client

    try:
        _b6_mod.httpx.Client = _FakeClient
        _b6_c = ResendClient(api_key="test_re_key", default_from="me@example.com")

        # T1: success
        _b6_r = _b6_c.send_email(to="x@y.com", subject="hi", html="<p>hi</p>", from_email="me@example.com")
        test("B6: send_email returns id on success", lambda: _b6_r == {"id": "abc-123"})

        # T2: body shape
        _b6_b = _b6_last_body["val"]
        test("B6: send_email shapes request body correctly", lambda: (
            _b6_b is not None
            and _b6_b["from"] == "me@example.com"
            and _b6_b["to"] == ["x@y.com"]
            and _b6_b["subject"] == "hi"
            and _b6_b["html"] == "<p>hi</p>"
            and "reply_to" not in _b6_b
            and "text" not in _b6_b
        ))

        # T3: send_email_text
        _b6_last_body["val"] = None
        _b6_c.send_email_text(to=["a@b.com", "c@d.com"], subject="plain", body="Hello\nWorld", from_email="me@example.com")
        _b6_b2 = _b6_last_body["val"]
        test("B6: send_email_text converts plain text to HTML", lambda: (
            _b6_b2 is not None
            and "Hello" in _b6_b2["html"]
            and "<br>" in _b6_b2["html"]
            and _b6_b2["text"] == "Hello\nWorld"
            and _b6_b2["to"] == ["a@b.com", "c@d.com"]
        ))

        # T4: error handling
        _b6_mod.httpx.Client = _FakeErrorClient
        _b6_raised = None
        try:
            _b6_c.send_email(to="x@y.com", subject="x", html="<p>x</p>")
        except ResendError as e:
            _b6_raised = e
        test("B6: error response raises ResendError with detail", lambda: (
            _b6_raised is not None
            and "422" in str(_b6_raised)
            and "missing from field" in str(_b6_raised)
        ))

    finally:
        _b6_mod.httpx.Client = _b6_orig

    # B7 — CalComClient (simple inline assertions)
    from api.integrations.cal_com_client import CalComClient, CalComError, _parse_iso

    test("B7: _parse_iso handles Z suffix", lambda: _parse_iso("2026-05-04T09:00:00Z") is not None)
    test("B7: _parse_iso handles +00:00", lambda: _parse_iso("2026-05-04T09:00:00+00:00") is not None)
    test("B7: _parse_iso returns None for garbage", lambda: _parse_iso("not-a-date") is None)

    class _FakeGetResponse:
        def __init__(self, status_code, json_data):
            self.status_code = status_code
            self._json = json_data
            self.text = json.dumps(json_data)
        @property
        def is_success(self):
            return self.status_code < 400
        def json(self):
            return self._json

    class _B7FakeClient:
        def __init__(self, timeout=None):
            self.timeout = timeout
        def __enter__(self):
            return self
        def __exit__(self, *args):
            pass
        def get(self, url, *, params, headers=None):
            _b7_last_params["val"] = params
            _b7_last_url["val"] = url
            _b7_last_headers["val"] = headers
            return _FakeGetResponse(200, {
                "status": "success",
                "data": {
                    "2026-05-04": [
                        {"start": "2026-05-04T09:00:00+02:00"},
                        {"start": "2026-05-04T09:30:00+02:00"},
                    ]
                }
            })

    class _B7FakeErrorClient:
        def __init__(self, timeout=None):
            self.timeout = timeout
        def __enter__(self):
            return self
        def __exit__(self, *args):
            pass
        def get(self, url, *, params, headers=None):
            return _FakeGetResponse(401, {"error": "Invalid API key"})

    _b7_last_params = {"val": None}
    _b7_last_url = {"val": None}
    _b7_last_headers = {"val": None}

    _b7_mod = __import__("api.integrations.cal_com_client", fromlist=["httpx"])
    _b7_orig = _b7_mod.httpx.Client

    try:
        _b7_mod.httpx.Client = _B7FakeClient
        _b7_c = CalComClient(api_key="test_cal_key", username="charle", default_event_type_id=123)

        # T1: success — check parsed shape
        _b7_slots = _b7_c.get_slots(event_type_id=123, days=7)
        test("B7: get_slots returns dict keyed by date", lambda: isinstance(_b7_slots, dict))
        test("B7: get_slots has 2026-05-04 key", lambda: "2026-05-04" in _b7_slots)
        test("B7: get_slots returns 2 slots for that date", lambda: len(_b7_slots["2026-05-04"]) == 2)
        test("B7: each slot has start/end/booking_url keys", lambda: all(
            {"start", "end", "booking_url"} <= set(s) for s in _b7_slots["2026-05-04"]
        ))
        _b7_first = _b7_slots["2026-05-04"][0]
        test("B7: slot start is ISO formatted", lambda: "2026-05-04T09:00:00" in _b7_first["start"])
        test("B7: slot end is 30m after start", lambda: "09:30:00" in _b7_first["end"])
        test("B7: booking_url contains username + slug", lambda: "charle" in _b7_first["booking_url"] and "30min" in _b7_first["booking_url"])
        test("B7: booking_url contains date param", lambda: "date=2026-05-04" in _b7_first["booking_url"])
        test("B7: booking_url contains slot param", lambda: "slot=" in _b7_first["booking_url"])

        # T2: request params + headers check (v2 uses header-based auth)
        _b7_p = _b7_last_params["val"]
        _b7_h = _b7_last_headers["val"]
        test("B7: request has eventTypeId=123 in params", lambda: _b7_p is not None and _b7_p.get("eventTypeId") == "123")
        test("B7: request has cal-api-key in headers", lambda: _b7_h is not None and _b7_h.get("cal-api-key") == "test_cal_key")
        test("B7: request URL is correct v2", lambda: _b7_last_url["val"] == "https://api.cal.com/v2/slots")

        # T3: error handling
        _b7_mod.httpx.Client = _B7FakeErrorClient
        _b7_raised = None
        try:
            _b7_c.get_slots(event_type_id=123)
        except CalComError as e:
            _b7_raised = e
        test("B7: 401 raises CalComError with detail", lambda: (
            _b7_raised is not None
            and "401" in str(_b7_raised)
            and "Invalid API key" in str(_b7_raised)
        ))

        # T4: empty slots response
        class _B7FakeEmptyClient:
            def __init__(self, timeout=None):
                self.timeout = timeout
            def __enter__(self):
                return self
            def __exit__(self, *args):
                pass
            def get(self, url, *, params, headers=None):
                return _FakeGetResponse(200, {})
        _b7_mod.httpx.Client = _B7FakeEmptyClient
        _b7_empty = _b7_c.get_slots(event_type_id=123)
        test("B7: empty response returns empty dict", lambda: _b7_empty == {})

    finally:
        _b7_mod.httpx.Client = _b7_orig

    # O3 — PhaseGateEnforcer
    from api.workflows.o3_phase_gate import PhaseGateEnforcer
    g = PhaseGateEnforcer()
    d_fail = g.check("define", {"problem_statement": "", "sipoc": None})
    test("O3: empty define fails", lambda: not d_fail.passed)
    d_good = g.check("define", {"problem_statement": "High TTF", "sipoc": {"S": "x"}})
    test("O3: good define passes", lambda: d_good.passed)

    # B8 — Candidate Search routes
    _b8_registered = ("candidates" in sys.modules.get("api.routes.candidates", {}).__name__
                      if False else True)
    # Actually check that the module imports cleanly
    try:
        from api.routes.candidates import router as _b8_router
        _b8_import_ok = True
    except Exception:
        _b8_import_ok = False
    test("B8: candidates router imports cleanly", lambda: _b8_import_ok)

    from api.routes.scheduling import router as _b8_sched_router
    test("B8: scheduling router imports cleanly", lambda: _b8_sched_router is not None)

    # TestClient-based integration tests for the candidates routes
    from fastapi.testclient import TestClient
    from fastapi import FastAPI
    from api.agents.specialists.s2_rag import RAGAgent, RAGResult

    _b8_app = FastAPI()
    _b8_app.include_router(_b8_router, prefix="/candidates")

    # Mock supabase that returns a fixed candidate
    class _B8MockSupabase:
        def __init__(self):
            self._fixed_candidate = {
                "candidate_id": "cand-b8-test-001",
                "name": "Java Dev",
                "email": "java@example.com",
                "phone": "+27 82 123 4567",
                "skills_json": ["Java", "Spring", "SQL", "Docker", "Kubernetes", "AWS"],
                "experience_summary": "Senior Java developer with 8 years of experience building enterprise applications.",
                "cv_storage_path": "cvs/cand-b8-test-001.docx",
                "is_duplicate": False,
                "missing_fields_json": [],
                "confidential": False,
            }

        def table(self, name):
            _b8_last_table["val"] = name
            return _B8MockTable(self, name)

        def rpc(self, name, params=None):
            return _B8MockRPCResponse([])

        def storage(self):
            return _B8MockStorage()

    class _B8MockTable:
        def __init__(self, parent, name):
            self._parent = parent
            self._name = name
            self._filters = {}

        def select(self, cols):
            return self

        def eq(self, k, v):
            self._filters[k] = v
            return self

        def in_(self, k, vals):
            self._filters[k] = vals
            return self

        def execute(self):
            if self._name == "candidates":
                if "cv_storage_path" in str(self._filters.get("candidate_id", "")):
                    # CV path query — return just the path
                    return _B8MockResponse([{
                        "cv_storage_path": self._parent._fixed_candidate["cv_storage_path"]
                    }])
                if "name, email" in str(self._filters.get("select", "")):
                    # schedule name/email query
                    pass
                # Default: return full candidate
                return _B8MockResponse([self._parent._fixed_candidate])
            return _B8MockResponse([])

    class _B8MockResponse:
        def __init__(self, data):
            self.data = data

    class _B8MockRPCResponse:
        def __init__(self, data):
            self.data = data
        def execute(self):
            return self

    class _B8MockStorage:
        def from_(self, bucket):
            return self
        def create_signed_url(self, path, expires_in):
            return {"signedURL": f"https://storage.example.com/{path}?token=abc"}

    _b8_last_table = {"val": None}

    # Monkey-patch RAGAgent on the test app's module so the route uses our mock
    _orig_rag = RAGAgent

    class _B8MockRAGAgent:
        def __init__(self, supabase):
            pass
        def retrieve(self, query, top_k=5, corpus_filter=None):
            return RAGResult(
                query=query,
                chunks=[
                    {
                        "chunk_id": "chunk-001",
                        "corpus_name": "cvs",
                        "chunk_text": "Java developer with Spring Boot experience",
                        "similarity": 0.92,
                        "metadata": {"candidate_id": "cand-b8-test-001"},
                    }
                ],
                context_window="Java developer with Spring Boot experience",
            )

    import api.routes.candidates as _b8_mod
    _b8_mod.RAGAgent = _B8MockRAGAgent

    @_b8_app.middleware("http")
    async def _b8_inject_mocks(request, call_next):
        request.state.supabase = _B8MockSupabase()
        request.state.llm = None
        response = await call_next(request)
        return response

    _b8_client = TestClient(_b8_app)

    # Test POST /candidates/search
    _b8_resp = _b8_client.post("/candidates/search", json={"q": "java", "limit": 3})
    test("B8: /search returns 200", lambda: _b8_resp.status_code == 200)
    _b8_data = _b8_resp.json()
    test("B8: /search returns results list", lambda: isinstance(_b8_data.get("results"), list))
    test("B8: /search returns at least one candidate",
         lambda: len(_b8_data.get("results", [])) >= 1)
    if _b8_data.get("results"):
        _b8_first = _b8_data["results"][0]
        test("B8: candidate card has id", lambda: "id" in _b8_first)
        test("B8: candidate card has name", lambda: "name" in _b8_first)
        test("B8: candidate card has email", lambda: "email" in _b8_first)
        test("B8: candidate card has skills", lambda: isinstance(_b8_first.get("skills"), list))
        test("B8: candidate card has match_score", lambda: isinstance(_b8_first.get("match_score"), (int, float)))
        test("B8: match_score is high (0.92)",
             lambda: _b8_first.get("match_score", 0) > 0.9)

    # Test POST /candidates/{id}/schedule with empty slots → 400
    _b8_sched_resp = _b8_client.post(
        "/candidates/cand-b8-test-001/schedule",
        json={"slots": [], "message": None},
    )
    test("B8: empty slots returns 400", lambda: _b8_sched_resp.status_code == 400)

    # Restore
    _b8_mod.RAGAgent = _orig_rag

# ── Level 2: Supabase Integration ──────────────────────────────────────────

def level2_supabase():
    print(f"\n{'='*60}\n📦 LEVEL 2: Supabase Integration\n{'='*60}")
    try:
        from supabase import create_client
        supa = create_client(os.environ["SUPABASE_URL"], os.environ["SUPABASE_SERVICE_KEY"])
        resp = supa.table("industry_benchmarks").select("*").eq("region", "South Africa").execute()
        test(f"D2: benchmarks count = {len(resp.data)}", lambda: len(resp.data) >= 3)
        resp2 = supa.table("roles").select("*").execute()
        test(f"Roles count = {len(resp2.data)}", lambda: len(resp2.data) >= 3)
    except Exception as e:
        test(f"Supabase connection: {e}", lambda: False)

# ── Level 3: LiteLLM Router (needs API key) ────────────────────────────────

def level3_litellm():
    print(f"\n{'='*60}\n📦 LEVEL 3: LiteLLM Router\n{'='*60}")
    if not os.environ.get("DEEPSEEK_API_KEY", ""):
        print(f"  {ANSI['YELLOW']}⚠ SKIP: DEEPSEEK_API_KEY not set{ANSI['RESET']}")
        return
    from api.tools.t3_litellm_router import LiteLLMRouter
    r = LiteLLMRouter()
    content, log = r.route("tagging", [
        {"role": "system", "content": "Answer in one word: is this metric or semantic?"},
        {"role": "user", "content": "What is our time to fill for senior Java devs?"},
    ])
    test(f"LiteLLM response received (${log.cost_usd:.6f}, {log.duration_ms}ms)", lambda: len(content) > 0)
    test(f"Model used: {log.model_used}", lambda: log.model_used == "deepseek-chat")

# ── Level 4: FastAPI Server ────────────────────────────────────────────────

def level4_fastapi():
    print(f"\n{'='*60}\n📦 LEVEL 4: FastAPI HTTP (target: {_TARGET})\n{'='*60}")
    import httpx
    try:
        r = httpx.get(f"{BACKEND_BASE}/health", timeout=5)
        test("GET /health", lambda: r.status_code == 200 and r.json()["status"] == "ok")
        r2 = httpx.post(f"{BACKEND_BASE}/trigger/manual", json={}, timeout=5)
        test("POST /trigger/manual", lambda: r2.status_code == 200 and r2.json().get("session_id"))
        r3 = httpx.post(f"{BACKEND_BASE}/trigger/goal-review", timeout=5)
        test("POST /trigger/goal-review", lambda: r3.status_code == 200)
    except Exception as e:
        test(f"Server not running: {e}", lambda: False)

# ── Level 5: Full E2E Kaizen (needs Supabase + LLM keys) ──────────────────

def level5_e2e_kaizen():
    print(f"\n{'='*60}\n📦 LEVEL 5: Full E2E Kaizen Lifecycle\n{'='*60}")
    if not os.environ.get("SUPABASE_SERVICE_KEY", "") or not os.environ.get("DEEPSEEK_API_KEY", ""):
        print(f"  {ANSI['YELLOW']}⚠ SKIP: Needs SUPABASE_SERVICE_KEY + DEEPSEEK_API_KEY{ANSI['RESET']}")
        return
    import uuid
    from datetime import datetime
    from supabase import create_client
    from api.tools.t3_litellm_router import LiteLLMRouter
    from api.workflows.o2_meta_orchestrator import MetaOrchestrator
    supa = create_client(os.environ["SUPABASE_URL"], os.environ["SUPABASE_SERVICE_KEY"])
    llm = LiteLLMRouter(supabase_client=supa)
    orch = MetaOrchestrator(supa, llm)
    sid = str(uuid.uuid4())
    # Insert kaizen_sessions row first — FK requirement for agent_invocations
    try:
        supa.table("kaizen_sessions").insert({
            "session_id": sid,
            "trigger_type": "manual",
            "phase": "detection",
            "status": "running",
            "output_state": {},
            "created_at": datetime.utcnow().isoformat(),
        }).execute()
    except Exception as e:
        test(f"Create session row: {e}", lambda: False)
        return
    result = orch.run_full_kaizen(session_id=sid, role_title="Senior Java Developer")
    # Update session status on completion
    try:
        final_status = "completed" if result.phase in ("complete", "detection") else "failed"
        supa.table("kaizen_sessions").update({
            "status": final_status,
            "phase": result.phase,
        }).eq("session_id", sid).execute()
    except Exception:
        pass
    test("E2E completed", lambda: result.phase is not None and result.phase != "error")
    test("D1: internal metrics computed", lambda: result.detection is not None and "time_to_fill_days" in str(result.detection))
    if result.define:
        test("K1: Define output", lambda: len(result.define.get("problem_statement", "")) > 0)
    if result.analyse:
        test("K3: Root causes found", lambda: len(result.analyse.get("root_causes", [])) > 0)
    print(f"  Phase: {result.phase}")
    if result.detection:
        gap = result.detection.get("gap", {})
        print(f"  Kaizen required: {gap.get('kaizen_required')}")

# ── MCP Server Test ─────────────────────────────────────────────────────────

def test_mcp_server():
    print(f"\n{'='*60}\n📦 MCP Server Test\n{'='*60}")
    from mcp_server import AnalyticsMCPServer
    mcp = AnalyticsMCPServer()
    tools = mcp.list_tools()
    test(f"MCP: {len(tools)} tools registered", lambda: len(tools) >= 4)
    r = mcp.handle_mcp_request({
        "method": "tools/call",
        "params": {"name": "benchmark_comparison", "arguments": {"internal_value": 45.0, "benchmark_median": 35.0}}
    })
    test("MCP: benchmark tool returns delta_pct", lambda: "delta_pct" in str(r))

# ── Runner ──────────────────────────────────────────────────────────────────

def load_env():
    """Load .env file if present."""
    env_path = os.path.join(os.path.dirname(__file__), ".env")
    if os.path.exists(env_path):
        with open(env_path) as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    k, v = line.split("=", 1)
                    os.environ.setdefault(k, v)

if __name__ == "__main__":
    load_env()

    if _LEVEL in (0, 1): level1_unit()
    if _LEVEL in (0, 2): level2_supabase()
    if _LEVEL in (0, 3): level3_litellm()
    if _LEVEL in (0, 4): level4_fastapi()
    if _LEVEL in (0, 5): level5_e2e_kaizen()
    if _LEVEL in (0, 6): test_mcp_server()

    print(f"\n{'='*60}")
    print(f"RESULTS: {ANSI['GREEN']}{pass_count} passed{ANSI['RESET']}, {ANSI['RED']}{fail_count} failed{ANSI['RESET']}")
    print(f"{'='*60}")
    sys.exit(0 if fail_count == 0 else 1)
